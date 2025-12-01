"""
Patent Document Processor
Handles extraction and storage of patent claims and title of invention
Supports both local SQLite and Turso cloud database.
"""

import os
import re
import sqlite3
from docx import Document as DocxDocument
from typing import Tuple, List

# Try to import Turso connection
try:
    from turso_db import get_turso_connection, is_turso_enabled
    TURSO_AVAILABLE = True
except ImportError:
    TURSO_AVAILABLE = False


class PatentClaimsDatabase:
    """SQLite database for storing patent claims and title"""

    def __init__(self, db_path="patent_claims.db"):
        self.db_path = db_path
        self.use_turso = TURSO_AVAILABLE and is_turso_enabled()
        self.init_database()

    def get_connection(self):
        """Get database connection (Turso or local SQLite)"""
        if self.use_turso:
            return get_turso_connection("patent_claims")
        return sqlite3.connect(self.db_path)

    def init_database(self):
        """Initialize SQLite database with tables for claims and title"""
        conn = self.get_connection()
        cursor = conn.cursor()

        # Table for title of invention
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS title_of_invention (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                source_document TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # Table for claims
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS claims (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                claim_number INTEGER,
                claim_text TEXT NOT NULL,
                source_document TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        conn.commit()
        conn.close()
        print(f"✅ Patent claims database initialized: {self.db_path}")

    def save_title(self, title: str, source_document: str):
        """Save title of invention to database"""
        conn = self.get_connection()
        cursor = conn.cursor()

        # Clear previous titles (keep only latest)
        cursor.execute("DELETE FROM title_of_invention")

        # Insert new title
        cursor.execute(
            "INSERT INTO title_of_invention (title, source_document) VALUES (?, ?)",
            (title, source_document)
        )

        conn.commit()
        conn.close()
        print(f"✅ Saved title of invention: {title[:50]}...")

    def save_claims(self, claims: List[Tuple[int, str]], source_document: str):
        """Save claims to database"""
        conn = self.get_connection()
        cursor = conn.cursor()

        # Clear previous claims
        cursor.execute("DELETE FROM claims")

        # Insert new claims
        for claim_number, claim_text in claims:
            cursor.execute(
                "INSERT INTO claims (claim_number, claim_text, source_document) VALUES (?, ?, ?)",
                (claim_number, claim_text, source_document)
            )

        conn.commit()
        conn.close()
        print(f"✅ Saved {len(claims)} claims to database")

    def get_title(self) -> str:
        """Retrieve title of invention from database"""
        conn = self.get_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT title FROM title_of_invention ORDER BY created_at DESC LIMIT 1")
        result = cursor.fetchone()
        conn.close()

        return result[0] if result else None

    def get_all_claims(self) -> List[Tuple[int, str]]:
        """Retrieve all claims from database"""
        conn = self.get_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT claim_number, claim_text FROM claims ORDER BY claim_number")
        results = cursor.fetchall()
        conn.close()

        return results

    def get_independent_claim(self) -> str:
        """Retrieve independent claim (claim 1) from database"""
        conn = self.get_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT claim_text FROM claims WHERE claim_number = 1 ORDER BY created_at DESC LIMIT 1")
        result = cursor.fetchone()
        conn.close()

        return result[0] if result else None


def extract_title_of_invention(docx_path: str) -> str:
    """
    Extract title of invention from Word document.
    The title is typically centered, in all caps, before "CLAIMS" or "What is claimed is:"

    Pattern:
    - All uppercase text
    - Usually centered
    - Appears before "CLAIMS" or "What is claimed is:"
    """
    try:
        doc = DocxDocument(docx_path)

        # Look for title in first few paragraphs
        potential_titles = []

        for i, para in enumerate(doc.paragraphs[:20]):  # Check first 20 paragraphs
            text = para.text.strip()

            # Stop if we reach claims section
            if any(marker in text.upper() for marker in ["WHAT IS CLAIMED", "CLAIMS"]):
                break

            # Check if this could be the title
            # Title characteristics:
            # 1. All uppercase or mostly uppercase
            # 2. Not empty
            # 3. Not too long (usually < 200 chars)
            # 4. Centered alignment (if possible to check)

            if text and len(text) > 10 and len(text) < 200:
                # Check if mostly uppercase (at least 80% of letters)
                letters = [c for c in text if c.isalpha()]
                if letters:
                    uppercase_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
                    if uppercase_ratio > 0.8:
                        potential_titles.append((i, text))

        # Find the title: usually the last all-caps paragraph before claims
        if potential_titles:
            # Return the last candidate (closest to CLAIMS section)
            title = potential_titles[-1][1]
            print(f"Extracted title: {title}")
            return title

        # Fallback: look for specific patterns
        full_text = "\n".join([p.text for p in doc.paragraphs[:20]])

        # Try to find text before "CLAIMS" or "What is claimed"
        claims_markers = [
            r"CLAIMS",
            r"What is claimed",
            r"What is defined",
            r"We claim"
        ]

        for marker in claims_markers:
            match = re.search(rf"(.+?)\n\s*{marker}", full_text, re.IGNORECASE | re.DOTALL)
            if match:
                before_claims = match.group(1).strip()
                # Get last line that's in uppercase
                lines = before_claims.split('\n')
                for line in reversed(lines):
                    line = line.strip()
                    if line and line.isupper() and len(line) > 10:
                        print(f"Extracted title (pattern match): {line}")
                        return line

        print("⚠️  Could not extract title of invention")
        return "TITLE NOT FOUND"

    except Exception as e:
        print(f"Error extracting title: {e}")
        return "ERROR EXTRACTING TITLE"


def extract_claims_exact(docx_path: str) -> List[Tuple[int, str]]:
    """
    Extract claims from Word document, preserving exact formatting.
    Each claim is saved as-is with all punctuation, spacing, and formatting preserved.

    Returns: List of tuples (claim_number, claim_text)
    """
    try:
        doc = DocxDocument(docx_path)
        claims = []

        # Find where claims section starts
        claims_start_idx = None
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if any(marker in text.upper() for marker in ["WHAT IS CLAIMED", "CLAIMS", "WHAT IS DEFINED", "WE CLAIM"]):
                claims_start_idx = i
                break

        if claims_start_idx is None:
            print("⚠️  Could not find claims section")
            return []

        # Extract claims starting from claims section
        current_claim_number = None
        current_claim_text = []

        for para in doc.paragraphs[claims_start_idx + 1:]:
            text = para.text

            # Skip empty paragraphs
            if not text.strip():
                if current_claim_text:
                    # Empty line might be part of claim formatting
                    current_claim_text.append(text)
                continue

            # Check if this starts a new claim (looks for "1.", "2.", etc.)
            claim_number_match = re.match(r"^(\d+)\.\s+", text)

            if claim_number_match:
                # Save previous claim if exists
                if current_claim_number is not None and current_claim_text:
                    claim_full_text = ''.join(current_claim_text)
                    claims.append((current_claim_number, claim_full_text))

                # Start new claim
                current_claim_number = int(claim_number_match.group(1))
                # Keep the exact text including the number
                current_claim_text = [text]
            else:
                # Continuation of current claim
                if current_claim_number is not None:
                    # Preserve exact spacing and line breaks
                    current_claim_text.append('\n' + text)

        # Save last claim
        if current_claim_number is not None and current_claim_text:
            claim_full_text = ''.join(current_claim_text)
            claims.append((current_claim_number, claim_full_text))

        print(f"✅ Extracted {len(claims)} claims preserving exact formatting")
        return claims

    except Exception as e:
        print(f"Error extracting claims: {e}")
        import traceback
        traceback.print_exc()
        return []


def process_patent_document(docx_path: str, db: PatentClaimsDatabase) -> dict:
    """
    Process patent document (Input Document 2):
    1. Extract title of invention
    2. Extract all claims with exact formatting
    3. Save to non-vector database (SQLite)

    Returns: dict with processing results
    """
    try:
        source_doc = os.path.basename(docx_path)

        # Extract title
        title = extract_title_of_invention(docx_path)
        db.save_title(title, source_doc)

        # Extract claims
        claims = extract_claims_exact(docx_path)
        db.save_claims(claims, source_doc)

        return {
            "success": True,
            "title": title,
            "num_claims": len(claims),
            "source": source_doc
        }

    except Exception as e:
        print(f"Error processing patent document: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e)
        }
